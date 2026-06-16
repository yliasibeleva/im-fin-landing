from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Поля страницы
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

DARK_BLUE = RGBColor(0x1a, 0x3a, 0x5c)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE_BG = RGBColor(0xDC, 0xE8, 0xF7)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    # Левая синяя полоска через border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '1a3a5c')
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def para(text, italic=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if italic:
        for run in p.runs:
            run.italic = True
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, highlight_last=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Заголовок
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '1a3a5c')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    # Строки
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        is_last = (ri == len(rows) - 1) and highlight_last
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if is_last:
                set_cell_bg(cell, 'dce8f7')
                run.bold = True
                run.font.color.rgb = DARK_BLUE
            elif ri % 2 == 1:
                set_cell_bg(cell, 'f0f4fa')
    doc.add_paragraph()
    return t

def note_box(title, text, color='d4a017'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = p.add_run(title + ' ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

def price_box(label, value, note):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(label + '\n')
    r1.font.size = Pt(11)
    r2 = p.add_run(value + '\n')
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = DARK_BLUE
    r3 = p.add_run(note)
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'ccd5e0')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)

# ── ШАПКА ──────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ООО «Империя Финанс»')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ЭКОНОМИЧЕСКОЕ ОБОСНОВАНИЕ\nСТОИМОСТИ БУХГАЛТЕРСКОГО СОПРОВОЖДЕНИЯ')
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Для ознакомления собственника / руководителя')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('г. Орёл, июнь 2026 г.')
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

divider()

para('Настоящий документ содержит обоснование стоимости услуг бухгалтерского сопровождения, предлагаемых ООО «Империя Финанс». Цена сформирована с учётом фактической сложности и объёма работ, выявленных в ходе предварительного анализа учётной базы организации.')
para('Предварительный анализ показал, что учётная база содержит системные отклонения, накопленные с 2021 года: незакрытые счета, отрицательные остатки, ошибки по лизингу и взаиморасчётам. Именно недостаточный уровень сопровождения в предыдущие периоды привёл к текущему состоянию учёта — и потребовал проведения комплексного восстановления. Стоимость настоящего предложения отражает тот уровень квалификации, при котором подобная ситуация не повторится.')

# ── ЧАСТЬ I ────────────────────────────────────────────────────────────────────
heading1('Часть I. Сопровождение ООО на ОСН — 75 000 руб./мес.')

heading2('1. Объём документооборота — выше среднего')
para('Только в I квартале 2026 года обработано 146 документов по приходу и реализации — около 50 документов в месяц. Для сравнения: стандартное торговое предприятие генерирует 20–30 документов в месяц. Каждый документ требует проверки, привязки к договору, контрагенту, номенклатуре и счёту учёта.')

add_table(
    ['Показатель', 'Норма (торговля)', 'Ваша организация'],
    [
        ['Документов в месяц', '20–30 шт.', '~50 шт.'],
        ['Превышение', '—', 'в 1,5–2 раза'],
    ]
)

heading2('2. Мини-производство / комплектация — нестандартный учёт')
para('Организация не продаёт то, что покупает: приобретаются комплектующие, которые собираются в готовую продукцию. Это принципиальное усложнение по сравнению с обычной торговлей:')
bullet('Ведение счетов учёта с субсчетами комплектации (сч. 20, 41)')
bullet('Корректное формирование себестоимости готового продукта')
bullet('Раздельный учёт входящего сырья и готовой продукции')
bullet('Правильное отражение в НДС и налоге на прибыль')
para('Данный блок работ не входит в стандартную торговую схему и требует отдельной квалификации.')

heading2('3. Несколько лизинговых договоров')
para('В учёте организации — несколько лизинговых объектов. Это существенная нагрузка для одной компании. По каждому договору ежемесячно выполняется:')
bullet('Разноска графика платежей: основной долг / проценты / НДС')
bullet('Учёт по ФСБУ 25/2018 — предмет аренды на балансе, обязательство по аренде')
bullet('Сверка с лизингодателем')
bullet('Корректное отражение в налоговом учёте')
para('Каждый лизинговый объект × полноценный ежемесячный учёт = объём, сопоставимый с отдельной ставкой специалиста на рынке.')

heading2('4. Перенос базы с 1С 7.7 на 1С 8')
para('В рамках сопровождения осуществляется перенос учётной базы с платформы 1С 7.7 на 1С 8: перенос справочников, остатков и документов, контроль корректности конвертации данных. Это дополнительная нагрузка, выходящая за рамки стандартного сопровождения.')
para('После завершения переноса сложность учёта не снижается: специфика деятельности организации (комплектация, лизинги, ОСН) остаётся неизменной, а новая платформа потребует настройки и адаптации учётных процессов.')

heading2('5. Из чего складывается стоимость 75 000 руб./мес.')
add_table(
    ['Блок работы', 'Доля трудоёмкости'],
    [
        ['Стандартное ведение ОСН (выписки, первичка, НДС, прибыль, зарплата, отчётность)', '55%'],
        ['Комплектация / мини-производство', '15%'],
        ['Лизинговые договоры (ФСБУ 25/2018)', '15%'],
        ['Перенос базы 1С 7→8', '15%'],
        ['Итого', '100%'],
    ]
)

heading2('6. Рыночное сравнение')
add_table(
    ['Уровень сложности', 'Рыночная стоимость'],
    [
        ['ИП, УСН, без НДС, до 30 документов', '10 000 – 20 000 руб./мес.'],
        ['ООО, ОСН, торговля, до 50 документов', '35 000 – 55 000 руб./мес.'],
        ['ООО, ОСН, производство / комплектация', 'от 60 000 руб./мес.'],
        ['+ несколько лизингов, перенос базы 1С 7→8', 'от 70 000 руб./мес.'],
    ]
)

note_box(
    'Сравнение с штатным бухгалтером:',
    'содержание бухгалтера аналогичного уровня в штате обойдётся в 90 000–120 000 руб. в месяц (оклад + НДФЛ + страховые взносы + больничные + отпускные + рабочее место). Аутсорсинг — полная ответственность и экономия не менее 30%.',
    'd4a017'
)

note_box(
    'Почему не подходит бухгалтер за 40 000 руб.?',
    'Текущее состояние учёта — незакрытые счета с 2021 года, ошибки по лизингу, отрицательные остатки — является прямым следствием недостаточной квалификации предыдущего специалиста. Низкая цена сопровождения уже обошлась организации дороже: потребовалось комплексное восстановление учёта. Специалист уровня, необходимого для данной организации, не работает за 40 000 рублей.',
    'c0392b'
)

price_box(
    'Стоимость сопровождения ООО на ОСН',
    '75 000 руб./мес.',
    'Повышение стоимости — только после согласования с Заказчиком'
)

divider()

# ── ЧАСТЬ II ───────────────────────────────────────────────────────────────────
heading1('Часть II. Сопровождение ИП на АУСН — 15 000 руб./мес.')

heading2('1. Что такое АУСН и почему бухгалтер необходим')
para('АУСН — автоматизированная упрощённая система налогообложения. Несмотря на название «автоматизированная», бухгалтер не исключается из процесса: он контролирует корректность данных, которые банк передаёт в налоговую. Ошибка в разметке дохода или расхода ведёт к неверному налогу напрямую, без права на пересчёт задним числом.')

heading2('2. Состав работ')
add_table(
    ['Блок', 'Содержание работ'],
    [
        ['Контроль доходов и расходов', 'Ежемесячная сверка с уполномоченным банком, проверка разметки операций'],
        ['Налоговый контроль', 'Проверка расчёта налога АУСН, контроль налоговой нагрузки'],
        ['Кадровый учёт', 'Приём, увольнение, отпуска, больничные, трудовые договоры'],
        ['Расчёт заработной платы', 'Начисление, выплата, отражение в банке'],
        ['Консультационная поддержка', 'Оперативные ответы по налогам и кадровым вопросам'],
    ],
    highlight_last=False
)

heading2('3. Рыночное сравнение')
add_table(
    ['Ситуация', 'Стоимость'],
    [
        ['ИП, УСН, без сотрудников, минимум операций', '5 000 – 8 000 руб./мес.'],
        ['ИП, УСН/АУСН, есть сотрудники, кадровый учёт', '12 000 – 18 000 руб./мес.'],
        ['ИП на АУСН с сотрудниками (наш случай)', '15 000 руб./мес.'],
    ]
)

price_box(
    'Стоимость сопровождения ИП на АУСН',
    '15 000 руб./мес.',
    'Включает налоговый контроль, кадровый учёт и расчёт заработной платы · Начало: 01.06.2026'
)

divider()

# ── ИТОГО ──────────────────────────────────────────────────────────────────────
heading1('Итоговая стоимость услуг')

add_table(
    ['Услуга', 'Стоимость'],
    [
        ['Ежемесячное сопровождение ООО на ОСН', 'от 65 000 до 75 000 руб./мес.'],
        ['Ежемесячное сопровождение ИП на АУСН', '15 000 руб./мес.'],
        ['Итого в месяц', 'от 80 000 до 90 000 руб./мес.'],
    ]
)

para('Все указанные цены сформированы исходя из фактически выявленного объёма и сложности учёта. Стоимость сопровождения может быть пересмотрена только при существенном изменении объёма операций и исключительно с предварительным согласованием с Заказчиком.')

# ── ПОДПИСЬ ────────────────────────────────────────────────────────────────────
divider()

p = doc.add_paragraph()
r = p.add_run('С уважением,')
r.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('Генеральный директор ООО «Империя Финанс»')
r.bold = True
r.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('Сибилева Юлия Николаевна')
r.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run('\n_________________________          Дата: ___.___.2026')
r.font.size = Pt(11)

out = r'D:\РС\Визитка\buh-bot\static\justification.docx'
doc.save(out)
print('OK:', out)
